import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse des performances
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "discrimin": "Discrimination des Nombres",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

def performance_level(score, quantiles):
    """Détermine le niveau de performance basé sur les quantiles."""
    if score >= quantiles[0.75]:  # Meilleurs 25%
        return "Maîtrise"
    elif score >= quantiles[0.50]:  # Moyenne haute
        return "Développement"
    else:
        return "Émergent"

def show_performance_school(df):
    """Affiche l'analyse des performances par école."""
    
    st.markdown("### 🏫 Performance Individuelle par École")
    
    # Sélection des indicateurs
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Indicateurs (partie gauche):",
            options=list(score_columns.keys())[:len(score_columns)//2],
            default=list(score_columns.keys())[:len(score_columns)//2],
            format_func=lambda x: score_columns[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Indicateurs (partie droite):",
            options=list(score_columns.keys())[len(score_columns)//2:],
            default=list(score_columns.keys())[len(score_columns)//2:],
            format_func=lambda x: score_columns[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        try:
            # Calcul des moyennes par école
            mean_scores_by_school = df.groupby("school")[selected_columns].mean().round(2).reset_index()
            
            # Calcul des niveaux de performance
            df_analysis = df.copy()
            df_analysis["total_score"] = df_analysis[selected_columns].sum(axis=1)
            mean_total_by_school = df_analysis.groupby("school")["total_score"].mean()
            
            quantiles = mean_total_by_school.quantile([0.50, 0.75])
            df_analysis["performance_level"] = df_analysis["total_score"].apply(
                lambda x: performance_level(x, quantiles)
            )
            
            # Affichage des moyennes par école
            st.subheader("📊 Moyenne des scores par école")
            st.dataframe(
                mean_scores_by_school,
                hide_index=True,
                use_container_width=True
            )
            
            # Export des résultats
            col1, col2 = st.columns(2)
            
            # Export CSV
            with col1:
                csv = mean_scores_by_school.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "📥 Télécharger en CSV",
                    csv,
                    "performance_school.csv",
                    "text/csv",
                    key='download-performance-csv'
                )
            
            # Export Word
            with col2:
                if st.button("📄 Exporter en Word"):
                    doc = create_performance_word_report(
                        df_analysis, 
                        mean_scores_by_school, 
                        selected_columns,
                        df_analysis.groupby(["school", "performance_level"]).size().unstack(fill_value=0)
                    )
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            "📥 Télécharger le rapport Word",
                            docx,
                            "performance_school_report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
            
            # Affichage des distributions
            st.subheader("📈 Distribution des scores par école")
            
            # Création des box plots en grille
            for i in range(0, len(selected_columns), 2):
                col1, col2 = st.columns(2)
                
                with col1:
                    if i < len(selected_columns):
                        fig = px.box(
                            df_analysis,
                            x="school",
                            y=selected_columns[i],
                            color="school",
                            title=f"Distribution de {score_columns[selected_columns[i]]}",
                            labels={
                                "school": "École",
                                selected_columns[i]: score_columns[selected_columns[i]]
                            }
                        )
                        fig.update_layout(
                            showlegend=False,
                            height=400,
                            xaxis_tickangle=-45
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    if i + 1 < len(selected_columns):
                        fig = px.box(
                            df_analysis,
                            x="school",
                            y=selected_columns[i + 1],
                            color="school",
                            title=f"Distribution de {score_columns[selected_columns[i + 1]]}",
                            labels={
                                "school": "École",
                                selected_columns[i + 1]: score_columns[selected_columns[i + 1]]
                            }
                        )
                        fig.update_layout(
                            showlegend=False,
                            height=400,
                            xaxis_tickangle=-45
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
            # Graphique des niveaux de performance
            st.subheader("📉 Répartition des niveaux de performance par école")
            
            performance_summary = df_analysis.groupby(
                ["school", "performance_level"]
            ).size().unstack(fill_value=0).reset_index()
            
            fig = px.bar(
                performance_summary.melt(
                    id_vars=["school"],
                    var_name="Niveau",
                    value_name="Nombre d'élèves"
                ),
                x="school",
                y="Nombre d'élèves",
                color="Niveau",
                barmode="stack",
                title="Répartition des niveaux de performance par école"
            )
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig, use_container_width=True)
            
        except Exception as e:
            st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")
    
    else:
        st.warning("Veuillez sélectionner au moins un indicateur à analyser.")

def create_performance_word_report(df, mean_scores, selected_columns, performance_data):
    """Crée un rapport Word avec les résultats des performances par école."""
    doc = Document()
    doc.add_heading("Analyse : Performance Individuelle par École", level=1)
    
    # Tableau des moyennes générales
    doc.add_heading("Moyenne des scores par école", level=2)
    table = doc.add_table(rows=1, cols=len(mean_scores.columns))
    table.style = 'Table Grid'
    
    # En-têtes
    for i, column in enumerate(mean_scores.columns):
        table.rows[0].cells[i].text = score_columns.get(column, column)
    
    # Données
    for _, row in mean_scores.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Box plots par indicateur
    doc.add_heading("Distribution des scores par école", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            fig = px.box(
                df,
                x="school",
                y=col,
                color="school",
                title=f"Distribution de {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()  # Espace
    
    # Graphique des niveaux de performance
    doc.add_heading("Répartition des niveaux de performance", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        performance_fig = px.bar(
            performance_data.reset_index().melt(
                id_vars=["school"],
                var_name="Niveau",
                value_name="Nombre d'élèves"
            ),
            x="school",
            y="Nombre d'élèves",
            color="Niveau",
            barmode="stack",
            title="Répartition des niveaux de performance par école"
        )
        img_path = os.path.join(tmp_dir, "performance_levels.png")
        performance_fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    return doc