import streamlit as st
import pandas as pd
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "discrimin": "Discrimination des Nombres",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

def show_language_comparison(df):
    """Affiche la comparaison entre écoles anglophones et néerlandophones."""
    
    st.markdown("""
    ### 🇳🇱🇬🇧 Comparaison entre Écoles Anglophones et Néerlandophones
    
    🔍 **Objectif** : Comparer les performances des élèves en fonction de la langue d'enseignement.
    
    📌 **Questions analysées** :
    - Les élèves des écoles anglophones performent-ils différemment de ceux des écoles néerlandophones ?
    - Y a-t-il des matières où un groupe surpasse systématiquement l'autre ?
    - Les différences observées sont-elles statistiquement significatives ?
    """)
    
    try:
        # Préparation des données
        df_analysis = df.copy()
        df_english = df_analysis[df_analysis["language_teaching"] == "English"]
        df_dutch = df_analysis[df_analysis["language_teaching"] == "Dutch"]
        
        # 1. Moyennes par langue d'enseignement
        st.subheader("📊 Moyenne des scores par langue d'enseignement")
        
        mean_scores = df_analysis.groupby("language_teaching")[list(score_columns.keys())].mean().round(2)
        mean_scores_melted = mean_scores.reset_index().melt(
            id_vars=["language_teaching"],
            var_name="Tâche",
            value_name="Score moyen"
        )
        mean_scores_melted["Tâche"] = mean_scores_melted["Tâche"].map(score_columns)
        
        fig = px.bar(
            mean_scores_melted,
            x="Tâche",
            y="Score moyen",
            color="language_teaching",
            barmode="group",
            title="Scores moyens par langue d'enseignement",
            labels={
                "language_teaching": "Langue d'enseignement",
                "Score moyen": "Score moyen"
            }
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 2. Distribution des scores
        st.subheader("📈 Distribution des scores selon la langue d'enseignement")
        
        for i in range(0, len(list(score_columns.keys())[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(score_columns):
                    col = list(score_columns.keys())[i]
                    fig = px.box(
                        df_analysis,
                        x="language_teaching",
                        y=col,
                        color="language_teaching",
                        title=f"Distribution - {score_columns[col]}",
                        labels={
                            "language_teaching": "Langue d'enseignement",
                            col: score_columns[col]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(list(score_columns.keys())[:6]):
                    col = list(score_columns.keys())[i + 1]
                    fig = px.box(
                        df_analysis,
                        x="language_teaching",
                        y=col,
                        color="language_teaching",
                        title=f"Distribution - {score_columns[col]}",
                        labels={
                            "language_teaching": "Langue d'enseignement",
                            col: score_columns[col]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
        
        # 3. Tests statistiques
        st.subheader("📊 Tests Statistiques (Mann-Whitney)")
        
        results_data = []
        for col in list(score_columns.keys())[:6]:
            english_scores = df_english[col].dropna()
            dutch_scores = df_dutch[col].dropna()
            
            if len(english_scores) > 0 and len(dutch_scores) > 0:
                stat, p = stats.mannwhitneyu(
                    english_scores,
                    dutch_scores,
                    alternative='two-sided'
                )
                results_data.append({
                    "Variable": score_columns[col],
                    "Statistique": f"{stat:.3f}",
                    "p-value": f"{p:.5f}",
                    "Interprétation": "Différence significative" if p < 0.05 else "Pas de différence significative"
                })
        
        results_df = pd.DataFrame(results_data)
        st.dataframe(results_df, hide_index=True)
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = results_df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "comparaison_langues.csv",
                "text/csv",
                key='download-language-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_language_comparison_word_report(
                    df_analysis,
                    list(score_columns.keys())[:6],
                    results_df,
                    mean_scores_melted
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "comparaison_langues_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
        
        # 4. Synthèse
        st.subheader("📝 Synthèse des résultats")
        
        significant_diff = results_df["p-value"].apply(lambda x: float(x) < 0.05).sum()
        total_tests = len(results_df)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric(
                "Différences significatives",
                f"{significant_diff} sur {total_tests} tests"
            )
        
        with col2:
            diff_percent = (significant_diff / total_tests) * 100
            st.metric(
                "Pourcentage de différences significatives",
                f"{diff_percent:.1f}%"
            )
    
    except Exception as e:
        st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")

def create_language_comparison_word_report(df, score_cols, results_df, mean_scores_df):
    """Crée un rapport Word avec les résultats de la comparaison des langues."""
    doc = Document()
    doc.add_heading("Analyse : Comparaison entre Écoles Anglophones et Néerlandophones", level=1)
    
    # Description
    doc.add_paragraph(
        "Objectif : Comparer les performances des élèves en fonction de la langue "
        "d'enseignement pour identifier les différences potentielles entre les groupes."
    )
    
    # Moyennes par langue
    doc.add_heading("Moyennes par Langue d'Enseignement", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        fig = px.bar(
            mean_scores_df,
            x="Tâche",
            y="Score moyen",
            color="language_teaching",
            barmode="group"
        )
        img_path = os.path.join(tmp_dir, "means.png")
        fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    # Distribution par langue
    doc.add_heading("Distribution des Scores", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.box(
                df,
                x="language_teaching",
                y=col,
                color="language_teaching",
                title=f"Distribution - {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Tests statistiques
    doc.add_heading("Résultats des Tests Statistiques", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    for i, header in enumerate(["Variable", "Statistique", "p-value", "Interprétation"]):
        table.rows[0].cells[i].text = header
    
    # Données
    for _, row in results_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc