import streamlit as st
import pandas as pd
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "discrimin": "Discrimination des Nombres",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

def show_gender_effect(df):
    """Affiche l'analyse de l'effet du genre sur les performances."""
    
    st.markdown("""
    ### 🚻 Effet du Genre sur la Performance des Élèves
    
    🔍 **Objectif** : Comparer les performances des filles et des garçons dans les différentes compétences évaluées.
    
    📌 **Questions analysées** :
    - Les filles et les garçons performent-ils différemment en lecture et en mathématiques ?
    - Y a-t-il des matières où un genre surpasse systématiquement l'autre ?
    - Les différences observées sont-elles statistiquement significatives ?
    """)
    
    try:
        # Préparation des données
        df_analysis = df.copy()
        df_analysis["gender"] = df_analysis["stgender"].map({1: "Garçon", 0: "Fille"}).fillna("Inconnu")
        
        # 1. Moyenne des scores par genre
        st.subheader("📊 Moyenne des scores par genre")
        
        mean_scores_by_gender = df_analysis.groupby("gender")[list(score_columns.keys())].mean().round(2)
        mean_scores_melted = mean_scores_by_gender.reset_index().melt(
            id_vars=["gender"],
            var_name="Tâche",
            value_name="Score moyen"
        )
        mean_scores_melted["Tâche"] = mean_scores_melted["Tâche"].map(score_columns)
        
        fig = px.bar(
            mean_scores_melted,
            x="Tâche",
            y="Score moyen",
            color="gender",
            barmode="group",
            title="Moyenne des scores par genre et par tâche"
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 2. Distribution des scores
        st.subheader("📈 Distribution des scores selon le genre")
        
        for i in range(0, len(list(score_columns.keys())[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(score_columns):
                    col = list(score_columns.keys())[i]
                    fig = px.box(
                        df_analysis,
                        x="gender",
                        y=col,
                        color="gender",
                        title=f"Distribution - {score_columns[col]}",
                        labels={
                            "gender": "Genre",
                            col: score_columns[col]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(score_columns):
                    col = list(score_columns.keys())[i + 1]
                    fig = px.box(
                        df_analysis,
                        x="gender",
                        y=col,
                        color="gender",
                        title=f"Distribution - {score_columns[col]}",
                        labels={
                            "gender": "Genre",
                            col: score_columns[col]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
        
        # 3. Tests statistiques
        st.subheader("📊 Tests Statistiques (Mann-Whitney)")
        
        results_data = []
        for col in list(score_columns.keys())[:6]:
            boys_scores = df_analysis[df_analysis["gender"] == "Garçon"][col].dropna()
            girls_scores = df_analysis[df_analysis["gender"] == "Fille"][col].dropna()
            
            if len(boys_scores) > 0 and len(girls_scores) > 0:
                stat, p = stats.mannwhitneyu(boys_scores, girls_scores, alternative='two-sided')
                results_data.append({
                    "Variable": score_columns[col],
                    "Statistique": f"{stat:.3f}",
                    "p-value": f"{p:.5f}",
                    "Interprétation": "Différence significative" if p < 0.05 else "Pas de différence significative"
                })
        
        results_df = pd.DataFrame(results_data)
        st.dataframe(results_df, hide_index=True)
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = results_df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "effet_genre.csv",
                "text/csv",
                key='download-gender-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_gender_effect_word_report(
                    df_analysis,
                    list(score_columns.keys())[:6],
                    results_df,
                    mean_scores_melted
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "effet_genre_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
    
    except Exception as e:
        st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")

def create_gender_effect_word_report(df, score_cols, results_df, mean_scores_df):
    """Crée un rapport Word avec les résultats de l'analyse de l'effet du genre."""
    doc = Document()
    doc.add_heading("Analyse : Effet du Genre sur la Performance des Élèves", level=1)
    
    # Description
    doc.add_paragraph(
        "Objectif : Comparer les performances des filles et des garçons "
        "dans les différentes compétences évaluées."
    )
    
    # Moyennes par genre
    doc.add_heading("Moyenne des scores par genre", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        fig = px.bar(
            mean_scores_df,
            x="Tâche",
            y="Score moyen",
            color="gender",
            barmode="group",
            title="Moyenne des scores par genre et par tâche"
        )
        img_path = os.path.join(tmp_dir, "mean_scores.png")
        fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    # Distribution par genre
    doc.add_heading("Distribution des scores selon le genre", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.box(
                df,
                x="gender",
                y=col,
                color="gender",
                title=f"Distribution - {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_gender_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Résultats des tests statistiques
    doc.add_heading("Résultats des Tests Statistiques", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    for i, header in enumerate(["Variable", "Statistique", "p-value", "Interprétation"]):
        table.rows[0].cells[i].text = header
    
    # Données
    for _, row in results_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc