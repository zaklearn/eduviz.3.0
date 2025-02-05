import streamlit as st
import pandas as pd
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes d'intérêt
columns_of_interest = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "st_nb_miss_school": "Nombre d'Absences",
    "st_nb_beenlate_school": "Nombre de Retards",
    "ses": "Statut Socio-Économique",
    "home_support": "Soutien à la Maison"
}

def show_statistics(df):
    """Affiche les statistiques descriptives et les visualisations."""
    
    # Sélection des colonnes
    st.subheader("📊 Sélection des variables")
    
    # Diviser les colonnes en deux colonnes pour l'affichage
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Variables (partie gauche):",
            options=list(columns_of_interest.keys())[:len(columns_of_interest)//2],
            default=list(columns_of_interest.keys())[:len(columns_of_interest)//2],
            format_func=lambda x: columns_of_interest[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Variables (partie droite):",
            options=list(columns_of_interest.keys())[len(columns_of_interest)//2:],
            default=list(columns_of_interest.keys())[len(columns_of_interest)//2:],
            format_func=lambda x: columns_of_interest[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        # Calcul des statistiques descriptives
        df_filtered = df[selected_columns]
        stats_summary = df_filtered.describe().round(2)
        
        # Affichage du tableau des statistiques
        st.subheader("📋 Tableau des Statistiques Descriptives")
        st.dataframe(stats_summary)
        
        # Boutons d'export
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = stats_summary.to_csv(index=True).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "statistiques.csv",
                "text/csv",
                key='download-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_word_report(df_filtered, stats_summary, selected_columns)
                # Sauvegarde temporaire du document
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "analyse1_Statistiques_descriptives.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
        
        # Affichage des histogrammes
        st.subheader("📊 Distribution des Scores")
        
        # Création des histogrammes en grille
        for i in range(0, len(selected_columns), 2):
            col1, col2 = st.columns(2)
            
            # Premier histogramme
            with col1:
                if i < len(selected_columns):
                    fig = px.histogram(
                        df_filtered,
                        x=selected_columns[i],
                        nbins=20,
                        marginal="box",
                        opacity=0.7,
                        title=f"Distribution de {columns_of_interest[selected_columns[i]]}"
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            # Deuxième histogramme
            with col2:
                if i + 1 < len(selected_columns):
                    fig = px.histogram(
                        df_filtered,
                        x=selected_columns[i + 1],
                        nbins=20,
                        marginal="box",
                        opacity=0.7,
                        title=f"Distribution de {columns_of_interest[selected_columns[i + 1]]}"
                    )
                    st.plotly_chart(fig, use_container_width=True)
    
    else:
        st.warning("Veuillez sélectionner au moins une variable à analyser.")

def create_word_report(df_filtered, stats_summary, selected_columns):
    """Crée un rapport Word avec les statistiques et les graphiques."""
    doc = Document()
    doc.add_heading("Analyse 1 : Statistiques Descriptives", level=1)
    
    # Ajout du tableau des statistiques
    doc.add_heading("Tableau des Statistiques Descriptives", level=2)
    table = doc.add_table(rows=1, cols=len(stats_summary.columns) + 1)
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = "Statistique"
    for idx, col in enumerate(stats_summary.columns):
        header_cells[idx + 1].text = columns_of_interest[col]
    
    # Données
    for stat_name in stats_summary.index:
        row_cells = table.add_row().cells
        row_cells[0].text = stat_name
        for idx, value in enumerate(stats_summary.loc[stat_name]):
            row_cells[idx + 1].text = str(value)
    
    # Ajout des histogrammes
    doc.add_heading("Distribution des Scores", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            fig = px.histogram(
                df_filtered,
                x=col,
                nbins=20,
                marginal="box",
                opacity=0.7,
                title=f"Distribution de {columns_of_interest[col]}"
            )
            
            # Sauvegarde temporaire de l'image
            img_path = os.path.join(tmp_dir, f"{col}_histogram.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph("")  # Espace après chaque image
    
    return doc