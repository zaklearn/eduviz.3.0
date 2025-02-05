import streamlit as st
import pandas as pd
import plotly.figure_factory as ff
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse des corrélations
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension"
}

def show_correlation(df):
    """Affiche l'analyse des corrélations entre les tâches."""
    
    # En-tête et description
    st.markdown("""
    ### 📉 Corrélations entre les tâches
    
    🔍 **Objectif** : Identifier les relations entre les différentes tâches du test EGRA/EGMA.
    
    📌 **Interprétation** : 
    - Une corrélation proche de 1 indique une forte relation positive
    - Une corrélation proche de -1 indique une forte relation négative
    - Une corrélation proche de 0 indique une faible relation
    - Les corrélations > 0.5 ou < -0.5 sont considérées comme significatives
    """)
    
    # Calcul de la matrice de corrélation
    corr_matrix = df[list(score_columns.keys())].corr().round(2)
    
    # Création du heatmap
    st.subheader("📊 Matrice de Corrélation")
    fig = ff.create_annotated_heatmap(
        z=corr_matrix.values,
        x=list(score_columns.values()),  # Utilisation des noms complets
        y=list(score_columns.values()),  # Utilisation des noms complets
        colorscale="viridis",
        showscale=True
    )
    
    # Ajustement de la mise en page du heatmap
    fig.update_layout(
        height=600,
        xaxis={'side': 'bottom'},
        xaxis_tickangle=-45
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Identification des corrélations significatives
    strong_correlations = []
    for i in range(len(corr_matrix)):
        for j in range(i + 1, len(corr_matrix)):  # Évite les doublons et la diagonale
            correlation = corr_matrix.iloc[i, j]
            if abs(correlation) > 0.5:
                strong_correlations.append({
                    'Tâche 1': list(score_columns.values())[i],
                    'Tâche 2': list(score_columns.values())[j],
                    'Corrélation': correlation
                })
    
    # Affichage des corrélations significatives
    if strong_correlations:
        st.subheader("📋 Corrélations significatives (>|0.5|)")
        df_strong = pd.DataFrame(strong_correlations)
        
        # Tri par valeur absolue de corrélation
        df_strong['Abs_Corr'] = df_strong['Corrélation'].abs()
        df_strong = df_strong.sort_values('Abs_Corr', ascending=False)
        df_strong = df_strong.drop('Abs_Corr', axis=1)
        
        # Formatage des corrélations
        df_strong['Corrélation'] = df_strong['Corrélation'].map('{:.2f}'.format)
        
        st.dataframe(df_strong, hide_index=True)
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = df_strong.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "correlations.csv",
                "text/csv",
                key='download-correlations-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_correlation_word_report(fig, df_strong)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "correlations_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
    else:
        st.info("Aucune corrélation significative (>|0.5|) n'a été trouvée.")

def create_correlation_word_report(heatmap_fig, df_strong):
    """Crée un rapport Word avec la matrice de corrélation et les corrélations significatives."""
    doc = Document()
    doc.add_heading("Analyse : Corrélations entre les tâches", level=1)
    
    # Ajout de la description
    doc.add_paragraph("Objectif : Identifier les relations entre les différentes tâches du test EGRA/EGMA.")
    doc.add_paragraph("Interprétation :")
    doc.add_paragraph("- Une corrélation proche de 1 indique une forte relation positive")
    doc.add_paragraph("- Une corrélation proche de -1 indique une forte relation négative")
    doc.add_paragraph("- Une corrélation proche de 0 indique une faible relation")
    doc.add_paragraph("- Les corrélations > 0.5 ou < -0.5 sont considérées comme significatives")
    
    # Ajout de la matrice de corrélation
    doc.add_heading("Matrice de Corrélation", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        img_path = os.path.join(tmp_dir, "heatmap.png")
        heatmap_fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    # Ajout des corrélations significatives
    doc.add_heading("Corrélations significatives (>|0.5|)", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = "Tâche 1"
    header_cells[1].text = "Tâche 2"
    header_cells[2].text = "Corrélation"
    
    # Données
    for _, row in df_strong.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Tâche 1"])
        row_cells[1].text = str(row["Tâche 2"])
        row_cells[2].text = str(row["Corrélation"])
    
    return doc