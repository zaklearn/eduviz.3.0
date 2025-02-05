import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.figure_factory as ff
from docx import Document
from docx.shared import Inches
import tempfile
import os
import numpy as np
import statsmodels.api as sm

# Définition des colonnes
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "discrimin": "Discrimination des Nombres",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

contextual_vars = {
    "st_english_home": "Langue Anglaise à la Maison",
    "st_dutch_home": "Langue Néerlandaise à la Maison",
    "st_other_language": "Autre Langue à la Maison",
    "ses": "Statut Socio-Économique (SES)",
    "home_support": "Soutien Parental"
}

def map_language(row):
    """Détermine le groupe linguistique basé sur les réponses."""
    if row["st_english_home"] == "Always":
        return "English Always"
    elif row["st_dutch_home"] == "Always":
        return "Dutch Always"
    elif row["st_english_home"] in ["Frequently", "Sometimes"]:
        return f"English {row['st_english_home']}"
    elif row["st_dutch_home"] in ["Frequently", "Sometimes"]:
        return f"Dutch {row['st_dutch_home']}"
    elif row["st_other_language"] == "Yes":
        return "Other Language"
    else:
        return "Other"

def validate_columns(df, required_columns):
    """Vérifie la présence des colonnes requises."""
    return [col for col in required_columns if col in df.columns]

def show_contextual_factors(df):
    """Affiche l'analyse des facteurs contextuels."""
    
    st.markdown("### 🏡 Facteurs Contextuels Favorisant l'Apprentissage")
    
    # Sélection des indicateurs
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Indicateurs (partie gauche):",
            options=list(score_columns.keys())[:len(score_columns)//2],
            default=list(score_columns.keys())[:len(score_columns)//2],
            format_func=lambda x: score_columns[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Indicateurs (partie droite):",
            options=list(score_columns.keys())[len(score_columns)//2:],
            default=list(score_columns.keys())[len(score_columns)//2:],
            format_func=lambda x: score_columns[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        try:
            # Préparation des données
            df_analysis = df.copy()
            df_analysis["language_group"] = df_analysis.apply(map_language, axis=1)
            df_analysis["total_score"] = df_analysis[selected_columns].sum(axis=1)
            
            # 1. Corrélations avec SES
            st.subheader("🔄 Relation entre SES et Scores")
            
            for i in range(0, len(selected_columns), 2):
                col1, col2 = st.columns(2)
                
                with col1:
                    if i < len(selected_columns):
                        fig = px.scatter(
                            df_analysis,
                            x="ses",
                            y=selected_columns[i],
                            trendline="ols",
                            title=f"SES vs {score_columns[selected_columns[i]]}",
                            labels={
                                "ses": "Statut Socio-Économique",
                                selected_columns[i]: score_columns[selected_columns[i]]
                            }
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    if i + 1 < len(selected_columns):
                        fig = px.scatter(
                            df_analysis,
                            x="ses",
                            y=selected_columns[i + 1],
                            trendline="ols",
                            title=f"SES vs {score_columns[selected_columns[i + 1]]}",
                            labels={
                                "ses": "Statut Socio-Économique",
                                selected_columns[i + 1]: score_columns[selected_columns[i + 1]]
                            }
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
            # 2. Impact du soutien parental
            st.subheader("👪 Impact du Soutien Parental sur les Scores")
            
            for i in range(0, len(selected_columns), 2):
                col1, col2 = st.columns(2)
                
                with col1:
                    if i < len(selected_columns):
                        fig = px.box(
                            df_analysis,
                            x="home_support",
                            y=selected_columns[i],
                            color="home_support",
                            title=f"Distribution de {score_columns[selected_columns[i]]}",
                            labels={
                                "home_support": "Soutien Parental",
                                selected_columns[i]: score_columns[selected_columns[i]]
                            }
                        )
                        fig.update_layout(showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    if i + 1 < len(selected_columns):
                        fig = px.box(
                            df_analysis,
                            x="home_support",
                            y=selected_columns[i + 1],
                            color="home_support",
                            title=f"Distribution de {score_columns[selected_columns[i + 1]]}",
                            labels={
                                "home_support": "Soutien Parental",
                                selected_columns[i + 1]: score_columns[selected_columns[i + 1]]
                            }
                        )
                        fig.update_layout(showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)
            
            # 3. Impact de la langue
            st.subheader("🗣️ Impact de la Langue Parlée à la Maison")
            
            fig = px.box(
                df_analysis,
                x="language_group",
                y="total_score",
                color="language_group",
                title="Distribution des Scores Totaux par Groupe Linguistique",
                labels={
                    "language_group": "Langue Parlée",
                    "total_score": "Score Total"
                }
            )
            fig.update_layout(
                showlegend=False,
                xaxis_tickangle=-45,
                height=500
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # 4. Matrice de corrélation
            st.subheader("📊 Matrice de Corrélation")
            
            correlation_vars = ["ses", "home_support", "total_score"]
            correlation_labels = ["SES", "Soutien Parental", "Score Total"]
            
            correlation_matrix = df_analysis[correlation_vars].corr().round(3)
            
            fig = ff.create_annotated_heatmap(
                z=correlation_matrix.values,
                x=correlation_labels,
                y=correlation_labels,
                colorscale="Viridis"
            )
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
            
            # Export des résultats
            correlation_data = correlation_matrix.reset_index().melt(
                id_vars="index",
                var_name="Variable",
                value_name="Corrélation"
            )
            
            col1, col2 = st.columns(2)
            
            # Export CSV
            with col1:
                csv = correlation_data.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "📥 Télécharger en CSV",
                    csv,
                    "facteurs_contextuels.csv",
                    "text/csv",
                    key='download-contextual-csv'
                )
            
            # Export Word
            with col2:
                if st.button("📄 Exporter en Word"):
                    doc = create_contextual_factors_word_report(
                        df_analysis,
                        selected_columns,
                        correlation_matrix
                    )
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            "📥 Télécharger le rapport Word",
                            docx,
                            "facteurs_contextuels_report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
            
        except Exception as e:
            st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")
    
    else:
        st.warning("Veuillez sélectionner au moins un indicateur à analyser.")

def create_contextual_factors_word_report(df, selected_columns, correlation_matrix):
    """Crée un rapport Word avec les résultats de l'analyse des facteurs contextuels."""
    doc = Document()
    doc.add_heading("Analyse : Facteurs Contextuels Favorisant l'Apprentissage", level=1)
    
    # Section SES
    doc.add_heading("Relations avec le Statut Socio-Économique (SES)", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            fig = px.scatter(
                df,
                x="ses",
                y=col,
                trendline="ols",
                title=f"SES vs {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_ses_scatter.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Section Soutien Parental
    doc.add_heading("Impact du Soutien Parental", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            fig = px.box(
                df,
                x="home_support",
                y=col,
                color="home_support",
                title=f"Distribution de {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_support_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Matrice de corrélation
    doc.add_heading("Matrice de Corrélation", level=2)
    table = doc.add_table(rows=len(correlation_matrix) + 1, cols=len(correlation_matrix) + 1)
    table.style = 'Table Grid'
    
    # En-têtes
    headers = [""] + list(correlation_matrix.columns)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Données
    for i, (idx, row) in enumerate(correlation_matrix.iterrows()):
        table.rows[i + 1].cells[0].text = idx
        for j, value in enumerate(row):
            table.rows[i + 1].cells[j + 1].text = f"{value:.3f}"
    
    return doc