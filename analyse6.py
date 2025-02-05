import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse
egra_columns = ["clpm", "phoneme", "sound_word", "cwpm", "listening", "orf", "comprehension"]
egma_columns = ["number_id", "discrimin", "missing_number", "addition", "subtraction", "problems"]

def cronbach_alpha(data):
    """Calcule le coefficient alpha de Cronbach."""
    items = data.dropna(axis=0)
    if items.empty or items.shape[1] < 2:  # Vérification des données suffisantes
        return None
        
    item_variances = items.var(axis=0, ddof=1)
    total_variance = items.sum(axis=1).var(ddof=1)
    num_items = items.shape[1]
    
    if total_variance == 0:  # Éviter la division par zéro
        return None
        
    alpha = (num_items / (num_items - 1)) * (1 - sum(item_variances) / total_variance)
    return alpha

def get_reliability_color(alpha):
    """Retourne une couleur selon la valeur de l'alpha."""
    if alpha is None:
        return "gray"
    elif alpha >= 0.9:
        return "green"
    elif alpha >= 0.7:
        return "lightgreen"
    elif alpha >= 0.6:
        return "orange"
    else:
        return "red"

def interpret_alpha(alpha):
    """Interprète la valeur de l'alpha de Cronbach."""
    if alpha is None:
        return "Données insuffisantes"
    elif alpha >= 0.9:
        return "Très bonne fiabilité"
    elif alpha >= 0.7:
        return "Bonne fiabilité"
    elif alpha >= 0.6:
        return "Acceptable"
    else:
        return "Faible fiabilité (le test doit être amélioré)"

def show_cronbach(df):
    """Affiche l'analyse de la fiabilité des tests."""
    
    st.markdown("""
    ### 📈 Fiabilité des tests (Cronbach Alpha)
    
    🔍 **Objectif** : Évaluer la cohérence interne des tests EGRA English, EGRA Dutch et EGMA.
    
    ℹ️ Le coefficient alpha de Cronbach est calculé pour mesurer la fiabilité des tests.
    
    📌 **Interprétation** :
    """)
    
    # Affichage des seuils dans une grille colorée
    col1, col2, col3, col4 = st.columns(4)
    col1.markdown("""
    <div style='background-color: #90EE90; padding: 10px; border-radius: 5px;'>
    ✔️ α ≥ 0.9<br>Très bonne fiabilité
    </div>
    """, unsafe_allow_html=True)
    
    col2.markdown("""
    <div style='background-color: #98FB98; padding: 10px; border-radius: 5px;'>
    ✔️ 0.7 ≤ α < 0.9<br>Bonne fiabilité
    </div>
    """, unsafe_allow_html=True)
    
    col3.markdown("""
    <div style='background-color: #FFA500; padding: 10px; border-radius: 5px;'>
    ✔️ 0.6 ≤ α < 0.7<br>Acceptable
    </div>
    """, unsafe_allow_html=True)
    
    col4.markdown("""
    <div style='background-color: #FFB6C6; padding: 10px; border-radius: 5px;'>
    ❌ α < 0.6<br>Faible fiabilité
    </div>
    """, unsafe_allow_html=True)
    
    # Calcul des coefficients alpha
    df_egra_english = df[df["language_teaching"] == "English"]
    df_egra_dutch = df[df["language_teaching"] == "Dutch"]
    
    alpha_egra_english = cronbach_alpha(df_egra_english[egra_columns])
    alpha_egra_dutch = cronbach_alpha(df_egra_dutch[egra_columns])
    alpha_egma = cronbach_alpha(df[egma_columns])
    
    # Création du DataFrame des résultats
    alpha_data = {
        "Test": ["EGRA English", "EGRA Dutch", "EGMA Mathématiques"],
        "Cronbach Alpha": [alpha_egra_english, alpha_egra_dutch, alpha_egma],
        "Fiabilité": [
            interpret_alpha(alpha_egra_english),
            interpret_alpha(alpha_egra_dutch),
            interpret_alpha(alpha_egma)
        ]
    }
    
    alpha_df = pd.DataFrame(alpha_data)
    
    # Affichage des résultats
    st.subheader("📊 Résultats du coefficient de Cronbach Alpha")
    
    # Formatage des valeurs alpha
    alpha_df["Cronbach Alpha"] = alpha_df["Cronbach Alpha"].apply(
        lambda x: f"{x:.3f}" if x is not None else "N/A"
    )
    
    # Affichage stylisé du tableau
    st.dataframe(
        alpha_df,
        hide_index=True,
        column_config={
            "Test": st.column_config.Column("Test", width="medium"),
            "Cronbach Alpha": st.column_config.Column("Cronbach Alpha", width="small"),
            "Fiabilité": st.column_config.Column("Fiabilité", width="medium")
        }
    )
    
    # Export des résultats
    col1, col2 = st.columns(2)
    
    # Export CSV
    with col1:
        csv = alpha_df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            "📥 Télécharger en CSV",
            csv,
            "cronbach_alpha.csv",
            "text/csv",
            key='download-cronbach-csv'
        )
    
    # Export Word
    with col2:
        if st.button("📄 Exporter en Word"):
            doc = create_cronbach_word_report(alpha_df)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc.save(tmp.name)
                with open(tmp.name, 'rb') as f:
                    docx = f.read()
                st.download_button(
                    "📥 Télécharger le rapport Word",
                    docx,
                    "cronbach_alpha_report.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            os.unlink(tmp.name)

def create_cronbach_word_report(alpha_df):
    """Crée un rapport Word avec les résultats de l'analyse de fiabilité."""
    doc = Document()
    doc.add_heading("Analyse : Fiabilité des tests (Cronbach Alpha)", level=1)
    
    # Ajout de la description
    doc.add_paragraph("Objectif : Évaluer la cohérence interne des tests EGRA English, EGRA Dutch et EGMA.")
    doc.add_paragraph("Le coefficient alpha de Cronbach est utilisé pour mesurer la fiabilité des tests.")
    
    # Ajout de l'interprétation
    doc.add_heading("Interprétation", level=2)
    interpretation = doc.add_paragraph()
    interpretation.add_run("• α ≥ 0.9 → Très bonne fiabilité\n")
    interpretation.add_run("• 0.7 ≤ α < 0.9 → Bonne fiabilité\n")
    interpretation.add_run("• 0.6 ≤ α < 0.7 → Acceptable\n")
    interpretation.add_run("• α < 0.6 → Faible fiabilité (le test doit être amélioré)")
    
    # Ajout des résultats
    doc.add_heading("Résultats du coefficient de Cronbach Alpha", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    for i, column in enumerate(["Test", "Cronbach Alpha", "Fiabilité"]):
        header_cells[i].text = column
    
    # Données
    for _, row in alpha_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc