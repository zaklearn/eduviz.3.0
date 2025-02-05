import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse des scores zéro
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension"
}

def show_zero_scores(df):
    """Affiche l'analyse des scores zéro."""
    
    # Sélection des colonnes
    st.subheader("⚠️ Sélection des tâches")
    
    # Diviser les colonnes en deux colonnes pour l'affichage
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Tâches (partie gauche):",
            options=list(score_columns.keys())[:len(score_columns)//2],
            default=list(score_columns.keys())[:len(score_columns)//2],
            format_func=lambda x: score_columns[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Tâches (partie droite):",
            options=list(score_columns.keys())[len(score_columns)//2:],
            default=list(score_columns.keys())[len(score_columns)//2:],
            format_func=lambda x: score_columns[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        try:
            # Calcul du pourcentage de scores zéro
            zero_scores = (df[selected_columns] == 0).sum()
            total_students = len(df)
            percentage_zero = ((zero_scores / total_students) * 100).round(2)
            
            df_zero_scores = pd.DataFrame({
                "Tâche": [score_columns[col] for col in selected_columns],
                "Pourcentage de Scores Zéro": percentage_zero
            })
            
            # Affichage du tableau
            st.subheader("📋 Proportion d'élèves ayant un score nul")
            st.dataframe(df_zero_scores)
            
            # Boutons d'export
            col1, col2 = st.columns(2)
            
            # Export CSV
            with col1:
                csv = df_zero_scores.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "📥 Télécharger en CSV",
                    csv,
                    "scores_zero.csv",
                    "text/csv",
                    key='download-zero-scores-csv'
                )
            
            # Export Word
            with col2:
                if st.button("📄 Exporter en Word"):
                    doc = create_zero_scores_word_report(df_zero_scores)
                    # Sauvegarde temporaire du document
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            "📥 Télécharger le rapport Word",
                            docx,
                            "zero_scores_export.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
            
            # Création et affichage du graphique
            st.subheader("📊 Pourcentage d'élèves avec un score de zéro par tâche")
            fig = px.bar(
                df_zero_scores,
                x="Pourcentage de Scores Zéro",
                y="Tâche",
                orientation="h",
                text="Pourcentage de Scores Zéro",
                color="Tâche",
                color_discrete_sequence=px.colors.sequential.Viridis
            )
            
            fig.update_layout(
                title="Pourcentage d'élèves avec un score de zéro",
                xaxis_title="Pourcentage (%)",
                yaxis_title="Tâches",
                height=400 + (len(selected_columns) * 30)  # Ajustement dynamique de la hauteur
            )
            
            st.plotly_chart(fig, use_container_width=True)
            
        except Exception as e:
            st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")
    
    else:
        st.warning("Veuillez sélectionner au moins une tâche à analyser.")

def create_zero_scores_word_report(df_zero_scores):
    """Crée un rapport Word avec les résultats des scores zéro."""
    doc = Document()
    doc.add_heading("Analyse 2 : Scores Zéro", level=1)
    
    # Ajout du tableau
    doc.add_heading("Proportion d'élèves ayant un score nul", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    header_cells[0].text = "Tâche"
    header_cells[1].text = "Pourcentage de Scores Zéro"
    
    # Données
    for _, row in df_zero_scores.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Tâche"])
        row_cells[1].text = f"{row['Pourcentage de Scores Zéro']}%"
    
    # Ajout du graphique
    doc.add_heading("Visualisation des scores zéro", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        fig = px.bar(
            df_zero_scores,
            x="Pourcentage de Scores Zéro",
            y="Tâche",
            orientation="h",
            text="Pourcentage de Scores Zéro",
            color="Tâche",
            color_discrete_sequence=px.colors.sequential.Viridis
        )
        
        # Sauvegarde temporaire du graphique
        img_path = os.path.join(tmp_dir, "zero_scores_graph.png")
        fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    return doc