import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse de l'effet des langues
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension"
}

def map_language(row):
    """Détermine le groupe linguistique basé sur les réponses."""
    if row["st_english_home"] == "Always":
        return "English Always"
    elif row["st_dutch_home"] == "Always":
        return "Dutch Always"
    elif row["st_english_home"] in ["Frequently", "Sometimes"]:
        return f"English {row['st_english_home']}"
    elif row["st_dutch_home"] in ["Frequently", "Sometimes"]:
        return f"Dutch {row['st_dutch_home']}"
    elif row["st_other_language"] == "Yes":
        return "Other Language"
    else:
        return "Other"

def show_language_effect(df):
    """Affiche l'analyse de l'effet des langues parlées à la maison."""
    
    # Sélection des colonnes
    st.subheader("🗣️ Sélection des indicateurs")
    
    # Diviser les colonnes en deux colonnes pour l'affichage
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Indicateurs (partie gauche):",
            options=list(score_columns.keys())[:len(score_columns)//2],
            default=list(score_columns.keys())[:len(score_columns)//2],
            format_func=lambda x: score_columns[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Indicateurs (partie droite):",
            options=list(score_columns.keys())[len(score_columns)//2:],
            default=list(score_columns.keys())[len(score_columns)//2:],
            format_func=lambda x: score_columns[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        try:
            # Création du groupe linguistique
            df_analysis = df.copy()
            df_analysis["language_group"] = df_analysis.apply(map_language, axis=1)
            
            # Calcul des statistiques descriptives par groupe linguistique
            stats_by_language = df_analysis.groupby("language_group")[selected_columns].describe().round(2)
            stats_by_language.columns = ['_'.join(col).strip() for col in stats_by_language.columns]
            stats_by_language = stats_by_language.reset_index()
            
            # Affichage du tableau des statistiques
            st.subheader("📋 Statistiques Descriptives par Groupe Linguistique")
            st.dataframe(stats_by_language)
            
            # Boutons d'export
            col1, col2 = st.columns(2)
            
            # Export CSV
            with col1:
                csv = stats_by_language.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "📥 Télécharger en CSV",
                    csv,
                    "language_effect.csv",
                    "text/csv",
                    key='download-language-effect-csv'
                )
            
            # Export Word
            with col2:
                if st.button("📄 Exporter en Word"):
                    doc = create_language_effect_word_report(df_analysis, stats_by_language, selected_columns)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            "📥 Télécharger le rapport Word",
                            docx,
                            "language_effect_export.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
            
            # Affichage des graphiques
            st.subheader("📊 Distribution des Scores par Groupe Linguistique")
            
            # Création d'une grille de graphiques (2 colonnes)
            for i in range(0, len(selected_columns), 2):
                col1, col2 = st.columns(2)
                
                # Premier graphique
                with col1:
                    if i < len(selected_columns):
                        fig = create_language_boxplot(
                            df_analysis,
                            selected_columns[i],
                            score_columns[selected_columns[i]]
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                # Deuxième graphique
                with col2:
                    if i + 1 < len(selected_columns):
                        fig = create_language_boxplot(
                            df_analysis,
                            selected_columns[i + 1],
                            score_columns[selected_columns[i + 1]]
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
        except Exception as e:
            st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")
    
    else:
        st.warning("Veuillez sélectionner au moins un indicateur à analyser.")

def create_language_boxplot(df, column, title):
    """Crée un box plot pour un indicateur donné."""
    fig = px.box(
        df,
        x="language_group",
        y=column,
        color="language_group",
        title=f"Distribution de {title}",
        labels={
            "language_group": "Groupe Linguistique",
            column: title
        }
    )
    fig.update_layout(
        showlegend=False,
        height=400,
        template="plotly_white",
        xaxis_tickangle=-45
    )
    return fig

def create_language_effect_word_report(df, stats_df, selected_columns):
    """Crée un rapport Word avec les résultats des effets de langue."""
    doc = Document()
    doc.add_heading("Analyse : Effet des Langues Parlées à la Maison", level=1)
    
    # Pour chaque indicateur sélectionné
    for col in selected_columns:
        # Ajout du titre de la section
        doc.add_heading(f"Statistiques pour {score_columns[col]}", level=2)
        
        # Création du tableau pour cet indicateur
        filtered_cols = [c for c in stats_df.columns if c.startswith(f"{col}_") or c == "language_group"]
        if filtered_cols:
            table = doc.add_table(rows=1, cols=len(filtered_cols))
            table.style = 'Table Grid'
            
            # En-têtes
            for i, col_name in enumerate(filtered_cols):
                table.rows[0].cells[i].text = col_name
            
            # Données
            for _, row in stats_df[filtered_cols].iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        
        # Ajout du graphique
        doc.add_paragraph()  # Espace avant le graphique
        with tempfile.TemporaryDirectory() as tmp_dir:
            fig = create_language_boxplot(df, col, score_columns[col])
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
        
        doc.add_paragraph()  # Espace après le graphique
    
    return doc