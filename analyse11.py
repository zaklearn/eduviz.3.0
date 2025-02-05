import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.figure_factory as ff
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "discrimin": "Discrimination des Nombres",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

def show_ses_home_support(df):
    """Affiche l'analyse de l'impact du SES et du soutien parental."""
    
    st.markdown("""
    ### 💰 Impact du Niveau Socio-Économique (SES) et du Soutien à Domicile
    
    🔍 **Objectif** : Analyser l'influence du niveau socio-économique (SES) et du soutien parental 
    sur les performances scolaires.
    
    📌 **Questions analysées** :
    - Le statut socio-économique (SES) influence-t-il les résultats scolaires ?
    - Un plus grand soutien parental est-il associé à de meilleures performances ?
    - Les relations observées sont-elles statistiquement significatives ?
    """)
    
    try:
        # Préparation des données
        df_analysis = df.copy()
        existing_scores = [col for col in score_columns.keys() if col in df.columns]
        df_analysis["total_score"] = df_analysis[existing_scores].sum(axis=1)
        
        # 1. Relation SES et scores
        st.subheader("📊 Relation entre SES et les Scores")
        
        for i in range(0, len(existing_scores[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(existing_scores):
                    fig = px.scatter(
                        df_analysis,
                        x="ses",
                        y=existing_scores[i],
                        trendline="ols",
                        title=f"SES vs {score_columns[existing_scores[i]]}",
                        labels={
                            "ses": "Statut Socio-Économique (SES)",
                            existing_scores[i]: score_columns[existing_scores[i]]
                        }
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(existing_scores[:6]):
                    fig = px.scatter(
                        df_analysis,
                        x="ses",
                        y=existing_scores[i + 1],
                        trendline="ols",
                        title=f"SES vs {score_columns[existing_scores[i + 1]]}",
                        labels={
                            "ses": "Statut Socio-Économique (SES)",
                            existing_scores[i + 1]: score_columns[existing_scores[i + 1]]
                        }
                    )
                    st.plotly_chart(fig, use_container_width=True)
        
        # 2. Impact du soutien parental
        st.subheader("📉 Impact du Soutien Parental sur les Scores")
        
        for i in range(0, len(existing_scores[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(existing_scores):
                    fig = px.box(
                        df_analysis,
                        x="home_support",
                        y=existing_scores[i],
                        color="home_support",
                        title=f"Distribution de {score_columns[existing_scores[i]]}",
                        labels={
                            "home_support": "Soutien Parental",
                            existing_scores[i]: score_columns[existing_scores[i]]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(existing_scores[:6]):
                    fig = px.box(
                        df_analysis,
                        x="home_support",
                        y=existing_scores[i + 1],
                        color="home_support",
                        title=f"Distribution de {score_columns[existing_scores[i + 1]]}",
                        labels={
                            "home_support": "Soutien Parental",
                            existing_scores[i + 1]: score_columns[existing_scores[i + 1]]
                        }
                    )
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
        
        # 3. Matrice de corrélation
        st.subheader("📈 Corrélations entre Variables")
        
        correlation_matrix = df_analysis[["ses", "home_support", "total_score"]].corr().round(3)
        
        fig = ff.create_annotated_heatmap(
            z=correlation_matrix.values,
            x=["SES", "Soutien Parental", "Score Total"],
            y=["SES", "Soutien Parental", "Score Total"],
            colorscale="Viridis",
            showscale=True
        )
        fig.update_layout(height=400)
        st.plotly_chart(fig, use_container_width=True)
        
        # 4. Tests statistiques
        st.subheader("📊 Tests Statistiques (Corrélation de Spearman)")
        
        results_data = []
        for col in existing_scores[:6]:
            correlation, p_value = stats.spearmanr(df_analysis["ses"].dropna(), df_analysis[col].dropna())
            results_data.append({
                "Variable": score_columns[col],
                "Coefficient": f"{correlation:.3f}",
                "p-value": f"{p_value:.5f}",
                "Interprétation": "Corrélation significative" if p_value < 0.05 else "Pas de corrélation significative"
            })
        
        results_df = pd.DataFrame(results_data)
        st.dataframe(results_df, hide_index=True)
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = results_df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "ses_soutien.csv",
                "text/csv",
                key='download-ses-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_ses_support_word_report(
                    df_analysis,
                    existing_scores[:6],
                    results_df,
                    correlation_matrix
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "ses_soutien_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
    
    except Exception as e:
        st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")

def create_ses_support_word_report(df, score_cols, results_df, correlation_matrix):
    """Crée un rapport Word avec les résultats de l'analyse SES et soutien parental."""
    doc = Document()
    doc.add_heading("Analyse : Impact du SES et du Soutien Parental", level=1)
    
    # Description
    doc.add_paragraph(
        "Objectif : Analyser l'influence du niveau socio-économique (SES) et du soutien parental "
        "sur les performances scolaires."
    )
    
    # Relations SES
    doc.add_heading("Relations avec le Statut Socio-Économique", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.scatter(
                df,
                x="ses",
                y=col,
                trendline="ols",
                title=f"SES vs {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_ses_scatter.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Impact du soutien
    doc.add_heading("Impact du Soutien Parental", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.box(
                df,
                x="home_support",
                y=col,
                color="home_support",
                title=f"Distribution de {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_support_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Matrice de corrélation
    doc.add_heading("Matrice de Corrélation", level=2)
    table = doc.add_table(rows=len(correlation_matrix) + 1, cols=len(correlation_matrix) + 1)
    table.style = 'Table Grid'
    
    # En-têtes
    headers = [""] + ["SES", "Soutien Parental", "Score Total"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Données
    for i, (idx, row) in enumerate(correlation_matrix.iterrows()):
        table.rows[i + 1].cells[0].text = ["SES", "Soutien Parental", "Score Total"][i]
        for j, value in enumerate(row):
            table.rows[i + 1].cells[j + 1].text = f"{value:.3f}"
    
    # Résultats des tests
    doc.add_heading("Résultats des Tests Statistiques", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    for i, header in enumerate(["Variable", "Coefficient", "p-value", "Interprétation"]):
        table.rows[0].cells[i].text = header
    
    # Données
    for _, row in results_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc