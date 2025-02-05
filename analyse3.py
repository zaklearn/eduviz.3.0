import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Définition des colonnes pour l'analyse des comparaisons entre écoles
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension"
}

def show_school_comparison(df):
    """Affiche l'analyse des comparaisons entre écoles."""
    
    # Sélection des colonnes
    st.subheader("🏫 Sélection des indicateurs")
    
    # Diviser les colonnes en deux colonnes pour l'affichage
    col1, col2 = st.columns(2)
    
    with col1:
        selected_columns_left = st.multiselect(
            "Indicateurs (partie gauche):",
            options=list(score_columns.keys())[:len(score_columns)//2],
            default=list(score_columns.keys())[:len(score_columns)//2],
            format_func=lambda x: score_columns[x]
        )
    
    with col2:
        selected_columns_right = st.multiselect(
            "Indicateurs (partie droite):",
            options=list(score_columns.keys())[len(score_columns)//2:],
            default=list(score_columns.keys())[len(score_columns)//2:],
            format_func=lambda x: score_columns[x]
        )
    
    selected_columns = selected_columns_left + selected_columns_right
    
    if selected_columns:
        try:
            # Calcul des statistiques descriptives par école
            stats_by_school = df.groupby("school")[selected_columns].describe().round(2)
            stats_by_school.columns = ['_'.join(col).strip() for col in stats_by_school.columns]
            stats_by_school = stats_by_school.reset_index()
            
            # Affichage du tableau des statistiques
            st.subheader("📋 Statistiques Descriptives par École")
            st.dataframe(stats_by_school)
            
            # Boutons d'export
            col1, col2 = st.columns(2)
            
            # Export CSV
            with col1:
                csv = stats_by_school.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    "📥 Télécharger en CSV",
                    csv,
                    "school_comparison.csv",
                    "text/csv",
                    key='download-school-comparison-csv'
                )
            
            # Export Word
            with col2:
                if st.button("📄 Exporter en Word"):
                    doc = create_school_comparison_word_report(df, stats_by_school, selected_columns)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            "📥 Télécharger le rapport Word",
                            docx,
                            "school_comparison_export.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
            
            # Affichage des graphiques
            st.subheader("📊 Distribution des Scores par École")
            
            # Création d'une grille de graphiques (2 colonnes)
            for i in range(0, len(selected_columns), 2):
                col1, col2 = st.columns(2)
                
                # Premier graphique
                with col1:
                    if i < len(selected_columns):
                        fig = px.box(
                            df,
                            x="school",
                            y=selected_columns[i],
                            color="school",
                            title=f"Distribution de {score_columns[selected_columns[i]]}",
                            labels={
                                "school": "École",
                                selected_columns[i]: score_columns[selected_columns[i]]
                            }
                        )
                        fig.update_layout(
                            showlegend=False,
                            height=400,
                            template="plotly_white"
                        )
                        st.plotly_chart(fig, use_container_width=True)
                
                # Deuxième graphique
                with col2:
                    if i + 1 < len(selected_columns):
                        fig = px.box(
                            df,
                            x="school",
                            y=selected_columns[i + 1],
                            color="school",
                            title=f"Distribution de {score_columns[selected_columns[i + 1]]}",
                            labels={
                                "school": "École",
                                selected_columns[i + 1]: score_columns[selected_columns[i + 1]]
                            }
                        )
                        fig.update_layout(
                            showlegend=False,
                            height=400,
                            template="plotly_white"
                        )
                        st.plotly_chart(fig, use_container_width=True)
            
        except Exception as e:
            st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")
    
    else:
        st.warning("Veuillez sélectionner au moins un indicateur à analyser.")

def create_school_comparison_word_report(df, stats_df, selected_columns):
    """Crée un rapport Word avec les résultats des comparaisons entre écoles."""
    doc = Document()
    doc.add_heading("Analyse 3 : Comparaison entre Écoles", level=1)
    
    # Pour chaque indicateur sélectionné
    for col in selected_columns:
        # Ajout du titre de la section
        doc.add_heading(f"Statistiques pour {score_columns[col]}", level=2)
        
        # Création du tableau pour cet indicateur
        filtered_cols = [c for c in stats_df.columns if c.startswith(f"{col}_") or c == "school"]
        if filtered_cols:
            table = doc.add_table(rows=1, cols=len(filtered_cols))
            table.style = 'Table Grid'
            
            # En-têtes
            for i, col_name in enumerate(filtered_cols):
                table.rows[0].cells[i].text = col_name
            
            # Données
            for _, row in stats_df[filtered_cols].iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        
        # Ajout du graphique
        doc.add_paragraph()  # Espace avant le graphique
        with tempfile.TemporaryDirectory() as tmp_dir:
            fig = px.box(
                df,
                x="school",
                y=col,
                color="school",
                title=f"Distribution de {score_columns[col]}"
            )
            
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
        
        doc.add_paragraph()  # Espace après le graphique
    
    return doc