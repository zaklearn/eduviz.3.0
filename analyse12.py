import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Standards internationaux
international_benchmarks = {
    "clpm": {"standard": 60, "nom": "Lettres Correctes Par Minute"},
    "phoneme": {"standard": 8, "nom": "Phonème"},
    "sound_word": {"standard": 6, "nom": "Mot Lu Correctement"},
    "cwpm": {"standard": 50, "nom": "Mots Corrects Par Minute"},
    "listening": {"standard": 3, "nom": "Écoute"},
    "orf": {"standard": 55, "nom": "Fluidité de Lecture Orale"},
    "comprehension": {"standard": 4, "nom": "Compréhension"},
    "number_id": {"standard": 25, "nom": "Identification des Nombres"},
    "discrimin": {"standard": 8, "nom": "Discrimination des Nombres"},
    "missing_number": {"standard": 7, "nom": "Nombre Manquant"},
    "addition": {"standard": 8, "nom": "Addition"},
    "subtraction": {"standard": 7, "nom": "Soustraction"},
    "problems": {"standard": 4, "nom": "Résolution de Problèmes"}
}

def show_international_comparison(df):
    """Affiche la comparaison avec les standards internationaux."""
    
    st.markdown("""
    ### 🌍 Comparaison avec les Standards Internationaux
    
    🔍 **Objectif** : Comparer les performances des élèves de Sint Maarten avec les standards internationaux.
    
    📌 **Questions analysées** :
    - Les élèves de Sint Maarten atteignent-ils les standards internationaux ?
    - Quelles compétences sont en retard par rapport aux attentes mondiales ?
    - Quelles sont les forces et les faiblesses du système éducatif local ?
    """)
    
    try:
        # Préparation des données
        existing_scores = [col for col in international_benchmarks.keys() if col in df.columns]
        mean_scores = df[existing_scores].mean()
        
        comparison_data = []
        for col in existing_scores:
            comparison_data.append({
                "Compétence": international_benchmarks[col]["nom"],
                "Code": col,
                "Moyenne Sint Maarten": round(mean_scores[col], 2),
                "Standard International": international_benchmarks[col]["standard"],
                "Écart": round(mean_scores[col] - international_benchmarks[col]["standard"], 2)
            })
        
        comparison_df = pd.DataFrame(comparison_data)
        
        # 1. Graphique de comparaison
        st.subheader("📊 Comparaison avec les standards internationaux")
        
        fig = px.bar(
            comparison_df.melt(
                id_vars="Compétence",
                value_vars=["Moyenne Sint Maarten", "Standard International"],
                var_name="Type",
                value_name="Score"
            ),
            x="Compétence",
            y="Score",
            color="Type",
            barmode="group",
            title="Performances vs Standards Internationaux"
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 2. Graphique des écarts
        st.subheader("📈 Écarts par rapport aux standards")
        
        fig = px.bar(
            comparison_df,
            x="Compétence",
            y="Écart",
            color="Écart",
            title="Écarts de performance",
            labels={"Écart": "Différence avec le standard"},
            color_continuous_scale="RdBu",
            color_continuous_midpoint=0
        )
        fig.add_hline(
            y=0,
            line_dash="dash",
            line_color="black",
            annotation_text="Standard",
            annotation_position="bottom right"
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 3. Tableau détaillé
        st.subheader("📋 Détails des écarts")
        
        # Formatage conditionnel des écarts
        st.dataframe(
            comparison_df,
            hide_index=True,
            column_config={
                "Compétence": st.column_config.Column("Compétence", width="medium"),
                "Code": st.column_config.Column("Code", width="small"),
                "Moyenne Sint Maarten": st.column_config.NumberColumn(
                    "Moyenne Sint Maarten",
                    format="%.2f"
                ),
                "Standard International": st.column_config.NumberColumn(
                    "Standard International",
                    format="%.2f"
                ),
                "Écart": st.column_config.NumberColumn(
                    "Écart",
                    format="%.2f"
                ),
            }
        )
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = comparison_df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "comparaison_internationale.csv",
                "text/csv",
                key='download-comparison-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_international_comparison_word_report(comparison_df)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "comparaison_internationale_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
        
        # 4. Synthèse des résultats
        st.subheader("📝 Synthèse des résultats")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric(
                "Compétences au-dessus des standards",
                f"{len(comparison_df[comparison_df['Écart'] > 0])} sur {len(comparison_df)}"
            )
        
        with col2:
            st.metric(
                "Plus grand écart",
                f"{comparison_df['Écart'].abs().max():.2f}",
                f"({comparison_df.loc[comparison_df['Écart'].abs().idxmax(), 'Compétence']})"
            )
    
    except Exception as e:
        st.error(f"Une erreur s'est produite lors de l'analyse : {str(e)}")

def create_international_comparison_word_report(comparison_df):
    """Crée un rapport Word avec les résultats de la comparaison internationale."""
    doc = Document()
    doc.add_heading("Analyse : Comparaison avec les Standards Internationaux", level=1)
    
    # Description
    doc.add_paragraph(
        "Objectif : Comparer les performances des élèves de Sint Maarten avec les "
        "standards internationaux pour identifier les forces et les points d'amélioration."
    )
    
    # Graphique de comparaison
    doc.add_heading("Comparaison des Performances", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        fig = px.bar(
            comparison_df.melt(
                id_vars="Compétence",
                value_vars=["Moyenne Sint Maarten", "Standard International"],
                var_name="Type",
                value_name="Score"
            ),
            x="Compétence",
            y="Score",
            color="Type",
            barmode="group"
        )
        img_path = os.path.join(tmp_dir, "comparison.png")
        fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    # Graphique des écarts
    doc.add_heading("Écarts de Performance", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        fig = px.bar(
            comparison_df,
            x="Compétence",
            y="Écart",
            color="Écart",
            color_continuous_scale="RdBu",
            color_continuous_midpoint=0
        )
        fig.add_hline(y=0, line_dash="dash", line_color="black")
        img_path = os.path.join(tmp_dir, "gaps.png")
        fig.write_image(img_path)
        doc.add_picture(img_path, width=Inches(6))
    
    # Tableau détaillé
    doc.add_heading("Détails des Résultats", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # En-têtes
    headers = ["Compétence", "Code", "Moyenne Sint Maarten",
              "Standard International", "Écart"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Données
    for _, row in comparison_df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(headers):
            row_cells[i].text = str(row[col])
    
    return doc