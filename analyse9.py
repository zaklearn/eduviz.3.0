import streamlit as st
import pandas as pd
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Variables d'analyse
class_practices_vars = ["teacher_experience", "teacher_training", "teaching_method", "use_of_materials"]
score_columns = {
    "clpm": "Lettres Correctes Par Minute",
    "phoneme": "Phonème",
    "sound_word": "Mot Lu Correctement",
    "cwpm": "Mots Corrects Par Minute",
    "listening": "Écoute",
    "orf": "Fluidité de Lecture Orale",
    "comprehension": "Compréhension",
    "number_id": "Identification des Nombres",
    "quantity": "Quantité",
    "missing_number": "Nombre Manquant",
    "addition": "Addition",
    "subtraction": "Soustraction",
    "problems": "Résolution de Problèmes"
}

def show_classroom_practices(df):
    """Affiche l'analyse des pratiques en classe."""
    
    st.markdown("""
    ### 📚 Pratiques en Classe
    
    🔍 **Objectif** : Identifier les pratiques pédagogiques et leur impact sur les performances des élèves.
    
    📌 **Questions analysées** :
    - Les enseignants plus expérimentés obtiennent-ils de meilleurs résultats ?
    - L'impact de la formation des enseignants est-il significatif ?
    - Une méthode pédagogique est-elle plus efficace que les autres ?
    - L'utilisation de matériel pédagogique influence-t-elle la performance ?
    """)
    
    # Vérification des colonnes existantes
    existing_scores = [col for col in score_columns.keys() if col in df.columns]
    
    if existing_scores:
        # 1. Impact de l'expérience enseignant
        st.subheader("👨‍🏫 Impact de l'Expérience Enseignant")
        
        for i in range(0, len(existing_scores[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(existing_scores):
                    fig = px.box(
                        df,
                        x="teacher_experience",
                        y=existing_scores[i],
                        color="teacher_experience",
                        title=f"Impact sur {score_columns[existing_scores[i]]}",
                        labels={
                            "teacher_experience": "Expérience Enseignant",
                            existing_scores[i]: score_columns[existing_scores[i]]
                        }
                    )
                    fig.update_layout(showlegend=False, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(existing_scores[:6]):
                    fig = px.box(
                        df,
                        x="teacher_experience",
                        y=existing_scores[i + 1],
                        color="teacher_experience",
                        title=f"Impact sur {score_columns[existing_scores[i + 1]]}",
                        labels={
                            "teacher_experience": "Expérience Enseignant",
                            existing_scores[i + 1]: score_columns[existing_scores[i + 1]]
                        }
                    )
                    fig.update_layout(showlegend=False, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
        
        # 2. Impact de la formation enseignant
        st.subheader("📚 Impact de la Formation Enseignant")
        
        for i in range(0, len(existing_scores[:6]), 2):
            col1, col2 = st.columns(2)
            
            with col1:
                if i < len(existing_scores):
                    fig = px.box(
                        df,
                        x="teacher_training",
                        y=existing_scores[i],
                        color="teacher_training",
                        title=f"Impact sur {score_columns[existing_scores[i]]}",
                        labels={
                            "teacher_training": "Formation Enseignant",
                            existing_scores[i]: score_columns[existing_scores[i]]
                        }
                    )
                    fig.update_layout(showlegend=False, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
            
            with col2:
                if i + 1 < len(existing_scores[:6]):
                    fig = px.box(
                        df,
                        x="teacher_training",
                        y=existing_scores[i + 1],
                        color="teacher_training",
                        title=f"Impact sur {score_columns[existing_scores[i + 1]]}",
                        labels={
                            "teacher_training": "Formation Enseignant",
                            existing_scores[i + 1]: score_columns[existing_scores[i + 1]]
                        }
                    )
                    fig.update_layout(showlegend=False, xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
        
        # 3. Comparaison des méthodes pédagogiques
        st.subheader("📖 Comparaison des Méthodes Pédagogiques")
        
        fig = px.box(
            df,
            x="teaching_method",
            y="cwpm",
            color="teaching_method",
            title="Impact sur la Lecture Fluide",
            labels={
                "teaching_method": "Méthode d'Enseignement",
                "cwpm": "Mots Corrects par Minute (CWPM)"
            }
        )
        fig.update_layout(showlegend=False, xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 4. Impact du matériel pédagogique
        st.subheader("🎒 Impact de l'Utilisation du Matériel Pédagogique")
        
        fig = px.box(
            df,
            x="use_of_materials",
            y="cwpm",
            color="use_of_materials",
            title="Impact sur la Lecture Fluide",
            labels={
                "use_of_materials": "Utilisation de Matériel",
                "cwpm": "Mots Corrects par Minute (CWPM)"
            }
        )
        fig.update_layout(showlegend=False, xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
        
        # 5. Tests statistiques
        st.subheader("📊 Tests Statistiques (Kruskal-Wallis)")
        
        results_data = []
        for col in existing_scores[:6]:
            groups = [df[df["teaching_method"] == method][col].dropna() 
                     for method in df["teaching_method"].unique()]
            stat, p = stats.kruskal(*groups)
            results_data.append({
                "Variable": score_columns[col],
                "Statistique": f"{stat:.3f}",
                "p-value": f"{p:.5f}",
                "Interprétation": "Différence significative" if p < 0.05 else "Pas de différence significative"
            })
        
        results_df = pd.DataFrame(results_data)
        st.dataframe(results_df, hide_index=True)
        
        # Export des résultats
        col1, col2 = st.columns(2)
        
        # Export CSV
        with col1:
            csv = results_df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 Télécharger en CSV",
                csv,
                "pratiques_classe.csv",
                "text/csv",
                key='download-practices-csv'
            )
        
        # Export Word
        with col2:
            if st.button("📄 Exporter en Word"):
                doc = create_classroom_practices_word_report(
                    df,
                    existing_scores[:6],
                    results_df
                )
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        "📥 Télécharger le rapport Word",
                        docx,
                        "pratiques_classe_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
    
    else:
        st.warning("Les colonnes nécessaires pour l'analyse ne sont pas disponibles dans les données.")

def create_classroom_practices_word_report(df, score_cols, results_df):
    """Crée un rapport Word avec les résultats de l'analyse des pratiques en classe."""
    doc = Document()
    doc.add_heading("Analyse : Pratiques en Classe", level=1)
    
    # Description
    doc.add_paragraph("Objectif : Identifier les pratiques pédagogiques et leur impact sur les performances des élèves.")
    
    # Expérience enseignant
    doc.add_heading("Impact de l'Expérience Enseignant", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.box(
                df,
                x="teacher_experience",
                y=col,
                color="teacher_experience",
                title=f"Impact sur {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_exp_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Formation enseignant
    doc.add_heading("Impact de la Formation Enseignant", level=2)
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in score_cols:
            fig = px.box(
                df,
                x="teacher_training",
                y=col,
                color="teacher_training",
                title=f"Impact sur {score_columns[col]}"
            )
            img_path = os.path.join(tmp_dir, f"{col}_train_box.png")
            fig.write_image(img_path)
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Tests statistiques
    doc.add_heading("Résultats des Tests Statistiques", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    for i, header in enumerate(["Variable", "Statistique", "p-value", "Interprétation"]):
        table.rows[0].cells[i].text = header
    
    # Données
    for _, row in results_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    return doc